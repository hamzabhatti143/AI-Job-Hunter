"""STEP 1 — Parse resume file (PDF or DOCX) into clean text."""
import io
import re
import json
from agents import function_tool


def _clean(text: str) -> str:
    """Normalise whitespace and remove junk characters."""
    text = re.sub(r'\r\n|\r', '\n', text)
    # Collapse 3+ blank lines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove non-printable except newline/tab
    text = re.sub(r'[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]', ' ', text)
    # Collapse spaces on a single line
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


async def resume_parser_impl(resume_path: str) -> str:
    try:
        text = ""
        path_lower = resume_path.lower()

        if path_lower.endswith(".pdf"):
            # Primary: pdfminer (handles most PDFs)
            try:
                from pdfminer.high_level import extract_text_to_fp
                from pdfminer.layout import LAParams
                output = io.StringIO()
                with open(resume_path, "rb") as f:
                    extract_text_to_fp(
                        f, output,
                        laparams=LAParams(line_margin=0.5, char_margin=2.0),
                        output_type="text",
                        codec="utf-8",
                    )
                text = output.getvalue()
            except Exception:
                text = ""

            # Fallback: pypdf if pdfminer returns nothing
            if not text.strip():
                try:
                    import pypdf
                    with open(resume_path, "rb") as f:
                        reader = pypdf.PdfReader(f)
                        text = "\n".join(
                            page.extract_text() or "" for page in reader.pages
                        )
                except Exception:
                    pass

            # Fallback 2: pymupdf (fitz)
            if not text.strip():
                try:
                    import fitz  # pymupdf
                    doc = fitz.open(resume_path)
                    text = "\n".join(page.get_text() for page in doc)
                    doc.close()
                except Exception:
                    pass

        elif path_lower.endswith((".docx", ".doc")):
            try:
                import docx
                doc = docx.Document(resume_path)
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                # Also get text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                parts.append(cell.text)
                text = "\n".join(parts)
            except Exception as e:
                return json.dumps({"success": False, "error": f"DOCX parse error: {e}"})
        else:
            return json.dumps({"success": False, "error": f"Unsupported format: {resume_path}"})

        if not text.strip():
            return json.dumps({
                "success": False,
                "error": "Could not extract text. The PDF may be image-based (scanned). Please use a text-based PDF or DOCX."
            })

        return json.dumps({
            "success": True,
            "resume_text": _clean(text),
            "file_path": resume_path,
        })

    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})


@function_tool
async def resume_parser_tool(resume_path: str) -> str:
    """Parse a resume file (PDF or DOCX) and return clean text. Tries multiple PDF engines."""
    return await resume_parser_impl(resume_path=resume_path)
